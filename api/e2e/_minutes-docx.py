#!/usr/bin/env python3
"""
Build a Google-Docs-compatible DOCX from the visual minutes HTML.

Deliberately plain: standard headings, simple tables, inline images, short
captions, and callout boxes rendered as single-cell shaded tables. No floats,
no absolute positioning, no SVG — so it survives upload to Google Drive and
opens in Google Docs with the layout intact.

Figure annotations are pre-flattened into the images (see _minutes-flatten.mjs),
so nothing depends on HTML overlays.
"""
import re, os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = '/var/public/pickleplay/docs'
HTML = f'{DOCS}/minutes-2026-07-08-visual.html'
FLAT = f'{DOCS}/minutes-assets/flat'
OUT  = f'{DOCS}/PickleBallers_Updated_Minutes_July_8_2026.docx'

OLIVE      = RGBColor(0x44, 0x62, 0x0A)
OLIVE_SOFT = RGBColor(0x6B, 0x8C, 0x2A)
INK        = RGBColor(0x2C, 0x35, 0x24)
MUTED      = RGBColor(0x6C, 0x7A, 0x5C)
BODY_FONT  = 'Calibri'          # ships with Google Docs + Word
HEAD_FONT  = 'Calibri'

# callout palettes: (fill, border, title colour)
CALLOUT = {
    'plain': ('F1F5E6', 'C6D4A8', OLIVE),
    'amber': ('FFF4E5', 'E0B072', RGBColor(0x8A, 0x54, 0x10)),
    'red':   ('FDECEC', 'E2A8A2', RGBColor(0x9B, 0x2C, 0x20)),
    'scope': ('F7F4FC', 'C3B2E6', RGBColor(0x4A, 0x2E, 0x86)),
    'qs':    ('F1F5E6', 'C6D4A8', OLIVE),
}
LABEL_FILL = {   # caption status labels
    'lbl-cur': 'E4EBF5', 'lbl-post': 'DFF0DC', 'lbl-mtg': 'EAE3F7',
    'lbl-prog': 'FFEFD6', 'lbl-cmp': 'E0EFF1', 'lbl-gap': 'FFF1DC',
}


def shade(cell, hexfill):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def borders(cell, hexcolor, sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:color'), hexcolor)
        b.append(e)
    tcPr.append(b)


def set_cell_margins(cell, top=90, bottom=90, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, v in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def runs_from(node, para, base_size=10.5, color=INK, italic=False):
    """Walk inline HTML, preserving <b>/<i> and the caption label chips."""
    for el in node.children:
        if isinstance(el, str):
            txt = re.sub(r'\s+', ' ', el)
            if txt.strip() or txt == ' ':
                r = para.add_run(txt)
                r.font.size = Pt(base_size); r.font.name = BODY_FONT
                r.font.color.rgb = color; r.italic = italic
            continue
        name = getattr(el, 'name', None)
        cls = el.get('class', []) if hasattr(el, 'get') else []
        if name == 'span' and 'lbl' in cls:
            key = next((c for c in cls if c.startswith('lbl-')), None)
            r = para.add_run(el.get_text(strip=True).upper() + '  ')
            r.font.size = Pt(7.5); r.font.name = HEAD_FONT; r.bold = True
            r.font.color.rgb = RGBColor(0x28, 0x4C, 0x7A)
            if key:
                hl = OxmlElement('w:shd'); hl.set(qn('w:val'), 'clear')
                hl.set(qn('w:fill'), LABEL_FILL.get(key, 'EEEEEE'))
                r._element.get_or_add_rPr().append(hl)
            continue
        if name == 'span' and 'figtag' in cls:
            r = para.add_run(el.get_text(strip=True).upper() + '  ')
            r.font.size = Pt(7.5); r.font.name = HEAD_FONT; r.bold = True
            r.font.color.rgb = RGBColor(0x44, 0x62, 0x0A)
            continue
        if name in ('b', 'strong'):
            r = para.add_run(re.sub(r'\s+', ' ', el.get_text()))
            r.bold = True; r.font.size = Pt(base_size); r.font.name = BODY_FONT
            r.font.color.rgb = color; r.italic = italic
        elif name in ('i', 'em'):
            r = para.add_run(re.sub(r'\s+', ' ', el.get_text()))
            r.italic = True; r.font.size = Pt(base_size); r.font.name = BODY_FONT
            r.font.color.rgb = color
        elif name == 'br':
            para.add_run().add_break()
        else:
            runs_from(el, para, base_size, color, italic)


def add_callout(doc, kind, title, body_nodes):
    fill, bd, tcol = CALLOUT[kind]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill); borders(c, bd); set_cell_margins(c)
    c.text = ''
    first = True
    if title:
        p = c.paragraphs[0]
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(10.5); r.font.name = HEAD_FONT
        r.font.color.rgb = tcol
        p.space_after = Pt(3)
        first = False
    for node in body_nodes:
        if node.name == 'ul':
            for li in node.find_all('li', recursive=False):
                p = c.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.space_after = Pt(2)
                runs_from(li, p, 10)
            continue
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.space_after = Pt(3)
        runs_from(node, p, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_figure(doc, img_path, caption_node, width_in):
    if not os.path.exists(img_path):
        print('  !! missing', img_path); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(6); p.space_after = Pt(2)
    p.add_run().add_picture(img_path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.paragraph_format.left_indent = Inches(0.12)
    cap.space_after = Pt(10)
    runs_from(caption_node, cap, 8.5, MUTED)


def main():
    soup = BeautifulSoup(open(HTML).read(), 'html.parser')
    doc = Document()

    st = doc.styles['Normal']
    st.font.name = BODY_FONT; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    for lvl, sz in ((1, 20), (2, 15), (3, 12), (4, 11)):
        h = doc.styles[f'Heading {lvl}']
        h.font.name = HEAD_FONT; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = OLIVE if lvl <= 2 else OLIVE_SOFT
        h.paragraph_format.space_before = Pt(14 if lvl <= 2 else 10)
        h.paragraph_format.space_after = Pt(5)

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.85)
    sec.left_margin = sec.right_margin = Inches(0.85)

    # footer: no recipient line
    f = sec.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run('PickleBallers — Updated Meeting Minutes')
    fr.font.size = Pt(8); fr.font.name = BODY_FONT; fr.font.color.rgb = MUTED

    # ── cover ───────────────────────────────────────────────────────────
    for txt, sz, col, bold, sp in [
        ('PICKLEBALLERS', 13, OLIVE_SOFT, True, 6),
        ('Updated Minutes of the Meeting', 26, OLIVE, True, 8),
        ('Product Direction, Current Delivery Status, and Decisions Required', 13, INK, False, 22),
        ('Meeting held: Wednesday, 8 July 2026', 11.5, OLIVE, True, 3),
        ('Status reviewed after development work on 9–10 July 2026', 11, INK, False, 3),
        ('Prepared for the PickleBallers Team', 11, INK, False, 22),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(sp)
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.name = HEAD_FONT
        r.bold = bold; r.font.color.rgb = col
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('This document combines the Wednesday discussion with the current state of the application '
                  'as reviewed on 12 July 2026. Screenshots are taken from the running product during active '
                  'development; features still being shaped are shown as clearly labelled concepts.')
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = MUTED; r.font.name = BODY_FONT
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── body ────────────────────────────────────────────────────────────
    body = soup.body
    fig_i = 0
    for node in body.find_all(recursive=False):
        cls = node.get('class', [])

        if node.name == 'section':       # HTML cover — replaced above
            continue

        if node.name == 'h2':
            if 'pagebreak' in cls:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            doc.add_heading(node.get_text(' ', strip=True), level=2)

        elif node.name in ('h3', 'h4'):
            if 'pagebreak' in cls:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            txt = node.get_text(' ', strip=True).replace(' UPDATE', '')
            doc.add_heading(txt, level=3 if node.name == 'h3' else 4)

        elif node.name == 'p':
            p = doc.add_paragraph()
            if 'footnote' in cls:
                runs_from(node, p, 9, MUTED, italic=True)
            else:
                runs_from(node, p, 10.5)

        elif node.name == 'ul':
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                p.space_after = Pt(3)
                runs_from(li, p, 10.5)

        elif node.name == 'ol':
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                p.space_after = Pt(3)
                runs_from(li, p, 10.5)

        elif node.name == 'div' and 'callout' in cls:
            kind = 'amber' if 'amber' in cls else 'red' if 'red' in cls else 'plain'
            ttl = node.find('div', class_='ttl')
            title = ttl.get_text(' ', strip=True).replace(' UPDATE', '') if ttl else None
            kids = [k for k in node.find_all(['p', 'ul', 'div'], recursive=False)
                    if 'ttl' not in (k.get('class') or []) and 'legend' not in (k.get('class') or [])]
            add_callout(doc, kind, title, kids)
            legend = node.find('div', class_='legend')
            if legend:
                for sp in legend.find_all('span', class_='i'):
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.space_after = Pt(1)
                    runs_from(sp, p, 9)

        elif node.name == 'div' and 'scope' in cls:
            add_callout(doc, 'scope', None, [node])

        elif node.name == 'div' and 'qs' in cls:
            ttl = node.find('div', class_='ttl')
            kids = [k for k in node.find_all(['ul'], recursive=False)]
            add_callout(doc, 'qs', ttl.get_text(strip=True) if ttl else None, kids)

        elif node.name == 'table':
            rows = node.find_all('tr')
            ncol = max(len(r.find_all(['td', 'th'])) for r in rows)
            t = doc.add_table(rows=0, cols=ncol)
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for tr in rows:
                cells = tr.find_all(['td', 'th'])
                row = t.add_row()
                for i, td in enumerate(cells):
                    if i >= ncol: break
                    c = row.cells[i]
                    c.text = ''
                    set_cell_margins(c, 70, 70, 100, 100)
                    p = c.paragraphs[0]; p.space_after = Pt(0)
                    if td.name == 'th':
                        shade(c, '44620A')
                        r = p.add_run(td.get_text(' ', strip=True))
                        r.bold = True; r.font.size = Pt(9.5); r.font.name = HEAD_FONT
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        tdcls = td.get('class', [])
                        if 'kv' in (tr.parent.parent.get('class') or []) and i == 0:
                            shade(c, 'F1F5E6')
                        runs_from(td, p, 9.5)
                        # colour the status column
                        stat = td.find(class_=re.compile('^st-')) or (td if any(
                            x.startswith('st-') for x in tdcls) else None)
                        if stat is not None:
                            col = {'st-done': RGBColor(0x2E, 0x7D, 0x32),
                                   'st-part': RGBColor(0xB2, 0x6A, 0x00),
                                   'st-pend': RGBColor(0xB2, 0x6A, 0x00),
                                   'st-expl': RGBColor(0x7A, 0x5C, 0xB8)}
                            key = next((x for x in (stat.get('class') or []) if x.startswith('st-')), None)
                            if key in col:
                                for r in p.runs:
                                    r.font.color.rgb = col[key]; r.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif node.name == 'figure':
            shot = node.find('div', class_='shot')
            cap = node.find('figcaption')
            img = shot.find('img')
            wide = bool(img and img['src'].split('/')[-1].startswith(('d', 'x0')) and 'x05' not in img['src'] and 'x06' not in img['src'])
            fname = f"{fig_i:02d}-" + (img['src'].split('/')[-1].replace('.png', '') if img else 'diagram') + '.png'
            add_figure(doc, f'{FLAT}/{fname}', cap, 6.4 if (wide or not img) else 2.5)
            fig_i += 1

        elif node.name == 'div' and ('phone-pair' in cls or 'phone-trio' in cls or 'phone' in cls):
            figs = node.find_all('figure')
            for fg in figs:
                shot = fg.find('div', class_='shot')
                img = shot.find('img')
                fname = f"{fig_i:02d}-" + (img['src'].split('/')[-1].replace('.png', '') if img else 'diagram') + '.png'
                # phones stack vertically in the DOCX — safer than side-by-side floats
                add_figure(doc, f'{FLAT}/{fname}', fg.find('figcaption'), 2.6)
                fig_i += 1

    doc.save(OUT)
    print('figures placed:', fig_i)
    print('DOCX →', OUT)


if __name__ == '__main__':
    main()
